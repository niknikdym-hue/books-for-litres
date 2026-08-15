#!/usr/bin/env python3
from __future__ import annotations

import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT / "ne-ver-golosu"
MANUSCRIPT = BOOK / "manuscript"
DIST = BOOK / "dist"
DIST.mkdir(parents=True, exist_ok=True)

ORDER = [
    "00-PROLOG.md", "01-VOICE.md", "02-FACE.md", "03-MESSAGES.md",
    "04-EVIDENCE.md", "05-DECISION.md", "06-SECRECY.md", "07-KNOWLEDGE.md",
    "07A-TRUST-CHAIN.md",
    "08-EXIT-CHANNEL.md", "09-RISK.md", "10-FAMILY.md", "11-BUSINESS.md",
    "11A-WHO-IS-REAL.md",
    "12-DETECTORS-PROVENANCE.md", "13-DIGITAL-DOUBLE.md",
    "13A-PUBLIC-EVIDENCE.md",
    "14-LIARS-DIVIDEND.md", "15-CONCLUSION.md", "16-APPENDICES.md", "17-SOURCES.md",
]

NON_TOC_H2_FILES = {
    "07A-TRUST-CHAIN.md",
    "11A-WHO-IS-REAL.md",
    "13A-PUBLIC-EVIDENCE.md",
    "17-SOURCES.md",
}

BANNED = [
    "мы живём в эпоху", "мы живем в эпоху", "важно понимать", "следует отметить",
    "в современном мире", "страшно подумать", "зловещая технология",
    "никому нельзя верить", "искусственный интеллект навсегда изменил всё",
    "искусственный интеллект навсегда изменил все",
]
OPEN_MARKERS = ["TODO", "VERIFY", "SOURCE?", "FIXME", "TBD"]


def strip_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\[(.+?)\]\((https?://[^)]+)\)", r"\1 (\2)", s)
    return s.replace("→", "—").strip()


def add_runs(p, text: str):
    pattern = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        token = m.group(0)
        r = p.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        r.bold = token.startswith("**")
        r.italic = not token.startswith("**")
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def blocks(text: str):
    out, buf = [], []

    def flush():
        nonlocal buf
        if buf:
            t = " ".join(x.strip() for x in buf if x.strip()).strip()
            if t:
                out.append(("p", t))
            buf = []

    for raw in text.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flush(); continue
        h = re.match(r"^(#{1,4})\s+(.*)$", line)
        if h:
            flush(); out.append((f"h{len(h.group(1))}", h.group(2).strip())); continue
        if re.match(r"^---+$", line.strip()):
            flush(); continue
        if line.startswith("> "):
            flush(); out.append(("quote", line[2:].strip())); continue
        b = re.match(r"^[-*]\s+(.*)$", line)
        if b:
            flush(); out.append(("bullet", b.group(1).strip())); continue
        n = re.match(r"^(\d+)\.\s+(.*)$", line)
        if n:
            # Keep the author's explicit number as text. Word's built-in List Number
            # style may silently continue numbering across separate lists, which is
            # undesirable in a reflowable LitRes manuscript.
            flush(); out.append(("number", f"{n.group(1)}. {n.group(2).strip()}")); continue
        buf.append(line)
    flush()
    return out


def build_docx(texts: list[tuple[str, str]]) -> Path:
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in [("Heading 1", 18), ("Heading 2", 15)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Internal subsection headings are intentionally NOT Word Heading 3/4.
    # LitRes builds navigation from Heading styles; keeping internal subheads in a
    # custom paragraph style prevents an oversized generated table of contents.
    internal = styles.add_style("Internal Heading", WD_STYLE_TYPE.PARAGRAPH)
    internal.base_style = normal
    internal.font.name = "Times New Roman"
    internal.font.size = Pt(12)
    internal.font.color.rgb = RGBColor(0, 0, 0)
    internal.font.bold = True
    internal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    internal.paragraph_format.space_before = Pt(8)
    internal.paragraph_format.space_after = Pt(3)

    # LitRes upload file intentionally has no title page, author block, manual TOC,
    # page numbers, headers/footers, cover, annotation or forced page breaks.
    for filename, text in texts:
        for kind, raw in blocks(text):
            if kind.startswith("h"):
                level = min(int(kind[1]), 4)
                is_reader_toc = level <= 2 and not (level == 2 and filename in NON_TOC_H2_FILES)
                p = doc.add_paragraph(style=f"Heading {level}" if is_reader_toc else "Internal Heading")
            elif kind == "quote":
                p = doc.add_paragraph(style="Quote")
            elif kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
            elif kind == "number":
                # Number already comes from the manuscript; keep it as ordinary
                # text so separate lists restart exactly where the author wrote 1.
                p = doc.add_paragraph()
            else:
                p = doc.add_paragraph()
            add_runs(p, raw)

    doc.core_properties.author = ""
    doc.core_properties.title = ""
    doc.core_properties.subject = ""
    doc.core_properties.keywords = ""
    out = DIST / "ne-ver-golosu-litres.docx"
    doc.save(out)
    return out


def plain(text: str) -> str:
    text = re.sub(r"^#{1,4}\s+.*$", "", text, flags=re.M)
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"[*_`>#]", "", text)
    return re.sub(r"\s+", " ", text).strip()


def qa(texts: list[tuple[str, str]], docx: Path) -> str:
    rows, all_plain = [], []
    paragraphs, sentences = defaultdict(list), defaultdict(list)
    for name, text in texts:
        ptext = plain(text)
        chars = len(ptext)
        words = len(re.findall(r"[A-Za-zА-Яа-яЁё0-9]+(?:[-’'][A-Za-zА-Яа-яЁё0-9]+)?", ptext))
        rows.append((name, chars, words))
        all_plain.append(ptext)
        for kind, value in blocks(text):
            if kind == "p":
                norm = re.sub(r"\s+", " ", strip_inline(value).lower()).strip()
                if len(norm) >= 100:
                    paragraphs[norm].append(name)
        for sentence in re.split(r"(?<=[.!?])\s+", ptext):
            norm = re.sub(r"\s+", " ", sentence.lower()).strip()
            if len(norm) >= 90:
                sentences[norm].append(name)

    whole = "\n".join(all_plain)
    lower = whole.lower()
    chars = len(whole)
    words = len(re.findall(r"[A-Za-zА-Яа-яЁё0-9]+(?:[-’'][A-Za-zА-Яа-яЁё0-9]+)?", whole))
    duplicates = [(p, locs) for p, locs in paragraphs.items() if len(set(locs)) > 1]
    dup_sent = [(s, locs) for s, locs in sentences.items() if len(set(locs)) > 1]
    banned = {x: lower.count(x) for x in BANNED if lower.count(x)}
    markers = {x: whole.count(x) for x in OPEN_MARKERS if whole.count(x)}
    watch = ["не нужно", "нужно", "важно", "представьте", "это не", "поэтому",
             "на самом деле", "не потому", "именно поэтому", "стоит запомнить"]

    r = ["# QUALITY REPORT — «НЕ ВЕРЬ ГОЛОСУ»", "",
         "Автоматическая проверка дополняет, но не заменяет содержательную редактуру.", "",
         f"- Общий объём текста: **{chars:,} знаков с пробелами**".replace(",", " "),
         f"- Ориентировочный объём: **{chars/40000:.2f} авторского листа**",
         f"- Слов: **{words:,}**".replace(",", " "),
         f"- DOCX: `{docx.relative_to(ROOT)}`", "", "## Объём по файлам", "",
         "| Файл | Знаков | Слов |", "|---|---:|---:|"]
    r.extend(f"| {name} | {c} | {w} |" for name, c, w in rows)
    r += ["", "## Gate", "",
          f"- Незакрытые TODO/VERIFY/SOURCE markers: **{markers or 'нет'}**",
          f"- Запрещённые шаблонные формулы: **{banned or 'нет'}**",
          f"- Точные дубли длинных абзацев между файлами: **{len(duplicates)}**",
          f"- Точные дубли длинных предложений между файлами: **{len(dup_sent)}**",
          "", "## Частотные контрольные обороты", ""]
    r.extend(f"- `{x}`: {lower.count(x)}" for x in watch)
    if duplicates:
        r += ["", "## Дубли абзацев — проверить вручную"]
        r.extend(f"- {locs}: {p[:220]}…" for p, locs in duplicates[:20])
    if dup_sent:
        r += ["", "## Дубли предложений — проверить вручную"]
        r.extend(f"- {locs}: {s[:220]}…" for s, locs in dup_sent[:20])
    problems = []
    if markers: problems.append("есть незакрытые редакционные маркеры")
    if banned: problems.append("есть запрещённые шаблонные обороты")
    if duplicates: problems.append("есть точные дубли длинных абзацев")
    if chars < 150000: problems.append("объём ниже целевого порога финальной экспертной редакции 150 000 знаков")
    r += ["", "## Автоматический вердикт",
          ("**REVIEW REQUIRED:** " + "; ".join(problems) + ".") if problems else
          "**AUTOMATED GATE PASS.** Финальный статус всё равно требует ручного фактчека, литературной редактуры и render-QA DOCX."]
    return "\n".join(r) + "\n"


def main():
    texts = []
    missing = []
    for name in ORDER:
        path = MANUSCRIPT / name
        if not path.exists():
            missing.append(str(path))
        else:
            texts.append((name, path.read_text(encoding="utf-8")))
    if missing:
        raise SystemExit("Missing manuscript files:\n" + "\n".join(missing))
    docx = build_docx(texts)
    report = qa(texts, docx)
    (DIST / "quality-report.md").write_text(report, encoding="utf-8")
    manifest = {"manuscript_files": ORDER, "docx": str(docx.relative_to(ROOT)),
                "quality_report": str((DIST / "quality-report.md").relative_to(ROOT))}
    (DIST / "build-manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(report)

if __name__ == "__main__":
    main()
